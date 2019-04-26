import sys
import os
from docx import Document
from openpyxl import load_workbook

"""
Tested in Python 3.6.4
"""

input_excel_path = sys.argv[1]
output_docx = input_excel_path.replace('xlsx', 'docx')

# loading excel file
xl_workbook = load_workbook(input_excel_path)
# selecting proper tab
xl_sheet = xl_workbook['AptivToFord_Interface']

# picking up row 2 (with headers)
header_row = [c.value for c in xl_sheet['2']]
# Finding which column stores 'Signal Name',
# which will be used as chapter title
signal_name_idx = header_row.index('Signal Name')

# opening word document
document = Document()

for row in xl_sheet.iter_rows():
    # iterate through rows which exist and which are not 'header' row
    if row[signal_name_idx].value and row[signal_name_idx].value != 'Signal Name':

        # check for strikethrough text
        if row[signal_name_idx].font.strike:
            p = document.add_paragraph(row[signal_name_idx].value + ' DELETED',
                                       style='Normal')
        else:
            p = document.add_paragraph(row[signal_name_idx].value,
                                       style='Normal')
        pb = document.add_paragraph(style='Body Text')
        for one_cell in row:
            if header_row[one_cell.col_idx - 1] and header_row[one_cell.col_idx - 1] != 'Signal Name':
                pb.add_run(header_row[one_cell.col_idx - 1]).bold = True
                pb.add_run(' ' + str(one_cell.value) + '\n')
document.save(output_docx)
