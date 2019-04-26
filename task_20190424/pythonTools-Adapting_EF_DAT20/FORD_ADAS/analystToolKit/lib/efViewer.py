import os
import csv
import cv2
import operator
import pickle
import docx
import re
from docx.shared import Pt
from docx.enum.text import WD_COLOR_INDEX, WD_PARAGRAPH_ALIGNMENT
from datetime import datetime
from functools import partial, reduce
from copy import deepcopy
from PyQt5 import QtCore, QtGui, QtWidgets
from delphiTools3.base import loadmat
from delphiTools3 import vis2

import glob
import pandas as pd
import numpy as np
from time import strftime
from collections import defaultdict

from uis import customHeader
from uis import videoEnchancer
from uis import bindDialog
from uis import reportGenWindow
from uis import resimViewer
from lib.tableViewModel import TableViewModel, CustomDelegate


class EFViewerClass:
    efv_frameSignal = QtCore.pyqtSignal()
    efv_eventSignal = QtCore.pyqtSignal(bool)
    efv_saveSignal = QtCore.pyqtSignal()

    def setup_efv(self):
        """
        Setup attributes of GUI: connect buttons to functions, set buttons clickable, set variables which will be used
        later to theirs default values
        :return: None
        """

        self.efv_actionCompare.triggered.connect(self.efv_onCompare)
        self.efv_actionCompare.setEnabled(False)
        self.efv_is2ndWindow = False
        self.saveID = 0

        self.efv_loadReportBtn.clicked.connect(partial(self.loadFile, self.efv_loadReportLine, self.efv_loadReport))
        self.efv_loadFolderBtn.clicked.connect(partial(self.efv_getFakeReport, self.efv_loadReport))

        self.efv_nextFrameBtn.clicked.connect(partial(self.efv_nextFrame, 2, self.efv_frameEdit, self.efv_updateFrame))
        self.efv_nextFrame10ShortCut = QtWidgets.QShortcut(QtGui.QKeySequence.fromString('Ctrl+Right'), self)
        self.efv_nextFrame10ShortCut.activated.connect(partial(self.efv_nextFrame, 20, self.efv_frameEdit, self.efv_updateFrame))
        self.efv_prevFrameBtn.clicked.connect(partial(self.efv_prevFrame, 2, self.efv_frameEdit, self.efv_updateFrame))
        self.efv_prevFrame10ShortCut = QtWidgets.QShortcut(QtGui.QKeySequence.fromString('Ctrl+Left'), self)
        self.efv_prevFrame10ShortCut.activated.connect(partial(self.efv_prevFrame, 20, self.efv_frameEdit, self.efv_updateFrame))

        self.efv_frameEdit.returnPressed.connect(partial(self.efv_setFrame, self.efv_updateFrame))
        self.efv_centerFrameBind = QtWidgets.QShortcut(QtGui.QKeySequence.fromString('Space'), self)
        self.efv_centerFrameBind.activated.connect(self.efv_centerFrame)

        self.efv_saveFrame10ShortCut = QtWidgets.QShortcut(QtGui.QKeySequence.fromString('Ctrl+E'), self)
        self.efv_saveFrame10ShortCut.activated.connect(self.efv_incrementSaveID)

        self.efv_nextEventBtn.clicked.connect(self.efv_nextEvent)
        self.efv_prevEventBtn.clicked.connect(self.efv_prevEvent)
        self.efv_eventEdit.returnPressed.connect(self.efv_setEvent)

        self.efv_loadBox.stateChanged.connect(partial(self.efv_updateEvent, reload=True))

        self.efv_dataTable.setContextMenuPolicy(QtCore.Qt.CustomContextMenu)
        self.efv_dataTable.customContextMenuRequested.connect(self.efv_customTableMenuRequested)
        self.efv_customizeTableAction = QtWidgets.QAction('Customize table', self)
        self.efv_customizeTableAction.triggered.connect(self.efv_setCustomHeader)
        self.efv_bindMenuAction = QtWidgets.QAction('Open bind menu', self)
        self.efv_bindMenuAction.triggered.connect(self.efv_openBindMenu)
        self.efv_saveAction = QtWidgets.QAction('Save events', self)
        self.efv_saveAction.setCheckable(True)
        self.efv_genCsv = QtWidgets.QAction('Generate csv', self)
        self.efv_genCsv.triggered.connect(self.efv_reportCsv)
        self.efv_genReport = QtWidgets.QAction('Generate report', self)
        self.efv_genReport.triggered.connect(self.efv_reportGen)

        self.efv_videoView.setContextMenuPolicy(QtCore.Qt.CustomContextMenu)
        self.efv_videoView.customContextMenuRequested.connect(self.efv_customFrameMenuRequested)
        scene = QtWidgets.QGraphicsScene(self)
        self.efv_videoView.setScene(scene)
        self.efv_videoView.setDragMode(QtWidgets.QGraphicsView.ScrollHandDrag)
        self.efv_videoView.wheelEvent = self.viewWheelEvent

        # Setup right-click menu items
        self.efv_videoEnchAction = QtWidgets.QAction('Enchance video', self)
        self.efv_videoEnchAction.triggered.connect(self.efv_enchVideo)

        self.efv_aflLayoutAction = QtWidgets.QAction('AFL overlay', self)
        self.efv_hrsLayoutAction = QtWidgets.QAction('HRS overlay', self)
        self.efv_tsrLayoutAction = QtWidgets.QAction('TSR overlay', self)
        self.efv_tsrPlusLayoutAction = QtWidgets.QAction('TSR Plus overlay', self)
        self.efv_obj2dLayoutAction = QtWidgets.QAction('Obj 2D overlay', self)
        self.efv_obj3dLayoutAction = QtWidgets.QAction('Obj 3D overlay', self)
        self.efv_lksIndLayoutAction = QtWidgets.QAction('LKS - Individual model overlay', self)
        self.efv_lksParLayoutAction = QtWidgets.QAction('LKS - Parallel model overlay', self)
        self.efv_lksHspLayoutAction = QtWidgets.QAction('LKS - High Speed Pred overlay', self)
        self.efv_lksReLayoutAction = QtWidgets.QAction('LKS - Road Edge', self)
        self.efv_lksBarLayoutAction = QtWidgets.QAction('LKS - Barrier', self)
        self.efv_lksHostLayoutAction = QtWidgets.QAction('LKS - Host Lane overlay', self)
        self.efv_lksNextLLayoutAction = QtWidgets.QAction('LKS - Next Left Lane overlay', self)
        self.efv_lksNextRLayoutAction = QtWidgets.QAction('LKS - Next Right Lane overlay', self)
        self.efv_lksBordLayoutAction = QtWidgets.QAction('LKS - Border overlay', self)
        self.efv_fusLayoutAction = QtWidgets.QAction('Fusion Tracks overlay', self)
        self.efv_fusPedLayoutAction = QtWidgets.QAction('Fusion - peds overlay', self)
        self.efv_tselLayoutAction = QtWidgets.QAction('TSEL overlay', self)
        self.efv_tselPathLayoutAction = QtWidgets.QAction('TSEL - path overlay', self)
        self.efv_pcaLayoutAction = QtWidgets.QAction('PCA', self)
        self.efv_failSafeLayoutAction = QtWidgets.QAction('Fail Safe / Severity lvl info', self)

        for item in [self.efv_aflLayoutAction, self.efv_hrsLayoutAction, self.efv_tsrLayoutAction,
                     self.efv_tsrPlusLayoutAction, self.efv_obj2dLayoutAction, self.efv_obj3dLayoutAction,
                     self.efv_lksIndLayoutAction, self.efv_lksParLayoutAction, self.efv_lksHspLayoutAction,
                     self.efv_lksReLayoutAction, self.efv_lksBarLayoutAction, self.efv_lksHostLayoutAction,
                     self.efv_lksNextLLayoutAction, self.efv_lksNextRLayoutAction, self.efv_lksBordLayoutAction,
                     self.efv_fusLayoutAction, self.efv_fusPedLayoutAction, self.efv_tselLayoutAction,
                     self.efv_tselPathLayoutAction, self.efv_pcaLayoutAction, self.efv_failSafeLayoutAction]:

            item.setCheckable(True)
            item.triggered.connect(self.efv_updateFrame)

        self.efv_gtLayoutAction = QtWidgets.QAction('GT overlay', self)
        self.efv_gtLayoutAction.setCheckable(True)
        self.efv_gtLayoutAction.triggered.connect(self.efv_updateFrame)

        self.splitter.setSizes([60, 1])
        self.splitter.splitterMoved.connect(self.efv_updateFrame)
        self.splitter_2.setSizes([1, 600])

        # Set variables to default values
        self.efv_videoHandler = None
        self.efv_currentEvent = 0
        self.efv_currentFrame = 0
        self.efv_header = []
        self.efv_mat = {'__path__': ''}
        self.efv_overlays_activated = False
        self.efv_active_overlays = {}
        self.efv_videoEnch = EnchancerWidget(self)
        self.efv_bindMenu = BindDialog(self)

    def efv_onCompare(self):
        '''
        Call Resim Viewer to enable resim data comparison. Function also sets GUI window params.
        :return: None
        '''
        self.rv_window = ResimViewer(self)
        self.rv_window.setWindowFlags(QtCore.Qt.Window |
                                      QtCore.Qt.CustomizeWindowHint |
                                      QtCore.Qt.WindowCloseButtonHint |
                                      QtCore.Qt.WindowMaximizeButtonHint |
                                      QtCore.Qt.WindowMinimizeButtonHint)
        self.rv_window.resize(1260, 960)
        self.rv_window.show()

    def efv_onDestroy(self):
        '''
        Destory Resim Viewer windows if tooltab is changed.
        :return: None
        '''
        if self.rv_window:
            self.rv_window.close()
        else:
            return

    def efv_loadReport(self, quietLoad=False):
        with open(self.efv_loadReportLine.text()) as csvfile:
            reader = csv.reader(csvfile)
            for row in reader:
                keys = row
                break
        self.efv_events = []
        with open(self.efv_loadReportLine.text()) as csvfile:
            reader = csv.DictReader(csvfile, keys)
            for i, event in enumerate(reader):
                if i:
                    self.efv_events.append(event)
        if self.efv_events:
            if not quietLoad:
                self.efv_header = []

            self.efv_nextEventBtn.setEnabled(True)
            self.efv_prevEventBtn.setEnabled(True)
            self.efv_eventEdit.setEnabled(True)

            if not quietLoad:
                self.efv_currentEvent = 0
            self.efv_eventEdit.setText(str(self.efv_currentEvent + 1))
            self.efv_eventLabel.setText(f'/{len(self.efv_events)}')

        if quietLoad:
            self.efv_updateEvent(quietLoad)
        else:
            self.efv_loadFolder(self.efv_loadReportLine, self.efv_loadDataLine, self.efv_updateEvent)

    def efv_loadFolder(self, dest_report, dest_folder, func=None):
        """
        Load folder containing logs to be viewed. This function (unlike one in toolKit.py) chcecks other field for path.
        This allows to start looking for files in same folder as the .csv report
        :param dest_report: QLine object with path to report
        :param dest_folder: QLine object with path to logs
        :param func: Function that will be executed at the end of efv_loadFolder
        :return: None
        """
        # If EF_report is not loaded
        if dest_report.text() == '':
            path = os.path.expanduser('~')

        # If EF_report is loaded and folder with logs is not loaded or contains old path
        elif dest_report.text() != '' and dest_folder.text() != dest_report.text():
            path = os.path.abspath(dest_report.text())

        # If folder with logs is loaded (report path will be loaded as well - by design)
        else:
            path = os.path.abspath(dest_folder.text())
        root = QtWidgets.QFileDialog.getExistingDirectory(self, directory=path,
                                                          options=QtWidgets.QFileDialog.ShowDirsOnly)
        if not root:
            return
        dest_folder.setText(root)

        if self.efv_is2ndWindow:
            self.rv_window.close()

        # remove old reports to prevent duplicate data append
        files = self.efv_getReportFiles()
        for file in files:
            os.remove(file)
            print(file + " removed.")

        self.saveID = 0

        self.efv_actionCompare.setEnabled(True)
        if func:
            func()

    def efv_getFakeReport(self, func=None):
        """
        Generate .csv file with "fake" EF report. File will contain logNames (based on .mat files in selected folder)
        and a columns "eventIndex" filled with -1. When reading this file ef_Viewer will set current frame to first one
        :param func: Function that will be executed at the end of report generation
        :return: None
        """
        search_path = os.path.expanduser('~')
        path = QtWidgets.QFileDialog.getExistingDirectory(self, directory=search_path,
                                                          options=QtWidgets.QFileDialog.ShowDirsOnly)
        if not path:
            return
        dir_path = os.path.abspath(path)

        mat_list = glob.glob(path + '\*.mat')

        frame_nums = pd.Series(np.full(len(mat_list), -1)).astype('int')
        lognames = pd.Series(mat_list).apply(lambda x: os.path.basename(x)[:-4])

        fake_ef_report = pd.DataFrame({'logName': lognames, 'eventIndex': frame_nums})
        _curr_time = strftime("%d-%m-%Y_%H-%M-%S")
        _out_file_name = "efv_Report_" + _curr_time + '.csv'
        _report_save_path = os.path.join(dir_path, _out_file_name)

        fake_ef_report.to_csv(_report_save_path, index=False)
        self.efv_loadReportLine.setText(_report_save_path)
        if func:
            func()

    def efv_load_groundtruth(self, file_path):
        if not os.path.isfile(file_path):
            print('Cannot load ground truth file!')
            return

        with open(file_path, 'rb') as f:
            dat_data = pickle.load(f)

        parsed_data = {}
        for frame in dat_data.keys():
            gid = None
            new_data = {}
            for obj in dat_data[frame]:
                if obj['objType'] == 'general':
                    gid = obj['frameGId']
                else:
                    pixel_height = obj['cords'][1][0] - obj['cords'][0][0]
                    if obj['infoDict']['height'] != '-':
                        obj_height = float(obj['infoDict']['height'][:4]) * 1000
                    else:
                        obj_height = 1650
                    try:
                        field = (5.47 * obj_height * 960) / (pixel_height * 4) / 1000
                    except:
                        field = 0
                    obj['infoDict']['distance'] = field

                ID = obj['objID']
                new_data[ID] = obj

            if gid is not None:
                parsed_data[gid] = new_data
            else:
                print('Gid parsing failed for frame {}'.format(frame))

        self.efv_mat['groundtruth'] = parsed_data

    def efv_nextFrame(self, step, frameEdit, updateFrame):
        """
        Get next video frame. Function checks if frame is in log's grab_indexes, adds step value and sets text to
        efv_FrameEdit (field below video image). Then it calls self.efv_currentFrame to update frame image
        (with overlays, coloring etc.)
        :param step: Value which will be added to current frame
        :return: None - function calls self.efv_updateFrame() to generate frame
        """
        currentFrame = int(frameEdit.text())
        if currentFrame + step <= self.matUpperIndex:
            currentFrame += step
        elif currentFrame + step >= self.matLowerIndex > self.matUpperIndex:
            currentFrame = (currentFrame + step) % 65536
        else:
            return

        frameEdit.setText(str(currentFrame))
        updateFrame()

    def efv_prevFrame(self, step, frameEdit, updateFrame):
        """
        See efv_nextFrame
        :param step:
        :return:
        """
        currentFrame = int(frameEdit.text())
        if currentFrame > self.matLowerIndex:
            currentFrame -= step
        elif self.matLowerIndex > self.matUpperIndex >= currentFrame:
            currentFrame = (currentFrame - step) % 65536
        else:
            return

        frameEdit.setText(str(currentFrame))
        updateFrame()

    @QtCore.pyqtSlot()
    def efv_updateFrameBySlider(self, frameEdit, updateFrame):
        """
        Function updates frame based on Video Slider position. Function updates current frame, sets text of
        frameEdit to new value and calls updateFrame()
        :return:
        """
        self.setFocus()
        videoSlider = self.sender()
        new_val = videoSlider.value()

        currentFrame = int(frameEdit.text())
        if new_val == currentFrame:
            return  # Prevents from multiple loops when updating slider
        if new_val in [self.matUpperIndex, self.matLowerIndex]:
            return   # bugfix: slider sets to min or max pos when new event is loaded
        if new_val % 2 == 0:
            new_val += 1

        frameEdit.setText(str(new_val))
        updateFrame()

    @QtCore.pyqtSlot()
    def efv_setFrame(self, updateFrame):
        self.setFocus()
        frameEdit = self.sender()

        currentFrame = int(frameEdit.text())
        if currentFrame % 2 == 0:
            currentFrame += 1
        if currentFrame > self.matUpperIndex and self.matLowerIndex < self.matUpperIndex:
            currentFrame = self.matUpperIndex
        elif currentFrame < self.matLowerIndex and self.matLowerIndex < self.matUpperIndex:
            currentFrame = self.matLowerIndex

        frameEdit.setText(str(currentFrame))
        updateFrame()

    def efv_centerFrame(self):
        self.efv_frameEdit.setText(self.efv_events[self.efv_currentEvent]['eventIndex'])
        self.efv_setFrame(self.efv_updateFrame)

    def _efv_initiate_overlays(self, videoFrame, mat, overlays_activated):
        '''
        Function initiates(creates instances of) overlay classes, stores them in dictionary
        :param videoFrame: video frame to activate overlays on
        :param mat: structure to get data from
        :param overlays_activated: flag to determine if overlays activation is needed
        :return: data to draw overlays on video frame
        '''
        if overlays_activated:
            return  # If overlays are already activated - do nothing

        frame_shape = videoFrame.shape
        # afl = vis2.AFLOverlay(mat, frame_shape)
        # hrs = vis2.HRSOverlay(mat, frame_shape)
        # tsr = vis2.TSROverlay(mat, frame_shape)
        # tsr_plus = vis2.TSRPlusOverlay(mat, frame_shape)
        obj = vis2.ObjectOverlay(mat, frame_shape)
        # lks = vis2.LKSOverlay(mat, frame_shape)
        # fs = vis2.FailSafeOverlay(mat, frame_shape)
        # fus = vis2.FusionOverlay(mat, frame_shape)
        tsel = vis2.TSELOverlay(mat, frame_shape)
        # tsel_path = vis2.TSELPathOverlay(mat, frame_shape)

        # overlays_data = {'AFL': afl, 'HRS': hrs, 'Tsr_plus': tsr_plus, 'TSR': tsr, 'Objects': obj,
        #                  'LKS': lks, 'Failsafe': fs, 'Fusion': fus, 'TSEL': tsel, 'Tsel_path': tsel_path}
        overlays_data = {'Objects': obj, 'TSEL': tsel}

        return overlays_data

    def _efv_draw_overlays(self, videoFrame, overlays_data, currentFrame):
        """
        Draw overlays on given frame. Function matches active overlays (checked by user) against instances of overlays
        in overlays_data (matches dictionaries keys names)
        :return: None - function draws overlays on frame_img in place
        """
        for overlay_name in self.efv_active_overlays.keys():
            if self.efv_active_overlays[overlay_name][0]:  # Check if given overlay is active - 1st element in list
                if len(self.efv_active_overlays[overlay_name]) > 1:  # If there are kwargs
                    draw_kwargs = self.efv_active_overlays[overlay_name][1].copy()  # Copy dictionary - create kwargs
                else:
                    draw_kwargs = {}

                for name, data in overlays_data.items():
                    if name in overlay_name:  # If active overlay's name matches an instance of overlay class
                        if 'data_dict' in draw_kwargs.keys():
                            # data_dict is used if one overlay can generate different outputs e.g. LKS can
                            # draw line or road edge
                            draw_kwargs['data_dict'] = getattr(data, draw_kwargs['data_dict'])
                        data.draw(videoFrame, currentFrame, **draw_kwargs)

    def efv_updateFrame(self):
        """
        Function updates frame.
        If a new event(log) is requested - self.efv_updateModel() function is called to get new log,
        else - function generates frame, checks which overlays are activated by user and calls self.efv_showVideoFrame
        to show generated frame_img with overlays and enhancement.
        In case of synchronization with resim viewer - function emits signal to update it.
        :return:
        """
        if self.efv_frameEdit.isEnabled():
            self.efv_currentFrame = int(self.efv_frameEdit.text())

            if not self.efv_updateModel(self.efv_dataTable, self.efv_mat, self.efv_currentFrame):
                return

            if self.efv_videoHandler is not None:
                self.efv_videoFrame = self.efv_generateVideoFrame(self.efv_videoView, self.efv_currentFrame)
                self.efv_active_overlays = {'AFL': [self.efv_aflLayoutAction.isChecked()],
                                            'HRS': [self.efv_hrsLayoutAction.isChecked()],
                                            'TSR': [self.efv_tsrLayoutAction.isChecked()],
                                            'Tsr_plus': [self.efv_tsrPlusLayoutAction.isChecked()],
                                            'Objects_2D': [self.efv_obj2dLayoutAction.isChecked(),
                                                           {'cube': False}],
                                            'Objects_3D': [self.efv_obj3dLayoutAction.isChecked(),
                                                           {'cube': True}],
                                            'LKS_ind': [self.efv_lksIndLayoutAction.isChecked(),
                                                        {'data_dict': 'data_ind', 'style': 'solid'}],
                                            'LKS_par': [self.efv_lksParLayoutAction.isChecked(),
                                                        {'data_dict': 'data_par', 'style': 'solid'}],
                                            'LKS_hsp': [self.efv_lksHspLayoutAction.isChecked(),
                                                        {'data_dict': 'data_hsp', 'style': 'solid'}],
                                            'LKS_RE': [self.efv_lksReLayoutAction.isChecked(),
                                                       {'data_dict': 'data_re', 'style': 're'}],
                                            'LKS_bar': [self.efv_lksBarLayoutAction.isChecked(),
                                                        {'data_dict': 'data_bar', 'style': 'bar'}],
                                            'LKS_host': [self.efv_lksHostLayoutAction.isChecked(),
                                                         {'data_dict': 'host_marker', 'style': 'solid'}],
                                            'LKS_nextL': [self.efv_lksNextLLayoutAction.isChecked(),
                                                          {'data_dict': 'next_left', 'style': 'solid'}],
                                            'LKS_nextR': [self.efv_lksNextRLayoutAction.isChecked(),
                                                          {'data_dict': 'next_right', 'style': 'solid'}],
                                            'LKS_bord': [self.efv_lksBordLayoutAction.isChecked(),
                                                         {'data_dict': 'border', 'style': 'solid'}],
                                            'Fusion_tracks': [self.efv_fusLayoutAction.isChecked(),
                                                              {'data_dict': 'veh_dict'}],
                                            'Fusion_peds': [self.efv_fusPedLayoutAction.isChecked(),
                                                            {'data_dict': 'ped_dict'}],
                                            'TSEL': [self.efv_tselLayoutAction.isChecked(),
                                                     {'data_dict': 'acc_moving_dict'}],
                                            'Tsel_path': [self.efv_tselPathLayoutAction.isChecked(),
                                                          {'data_dict': 'path_data'}],
                                            'TSEL_PCA': [self.efv_pcaLayoutAction.isChecked(),
                                                           {'data_dict': 'pca_moving_dict'}],
                                            'Failsafe': [self.efv_failSafeLayoutAction.isChecked()]
                                            }

                # Check if any overlay is checked by user
                if any(data[0] for data in list(self.efv_active_overlays.values())):
                    if not self.efv_overlays_activated:  # If overlays were not activated, activate them
                        self.efv_overlays_data = self._efv_initiate_overlays(self.efv_videoFrame, self.efv_mat, self.efv_overlays_activated)
                        # setting parameter overlays_activated to true will prevent overlays from being initiated every time a new frame is generated
                        self.efv_overlays_activated = True
                else:
                    self.efv_overlays_activated = False

            # signal emitted to resim viewer
            if self.efv_is2ndWindow:
                self.efv_frameSignal.emit()

            if self.efv_overlays_activated:  # If overlays are active, draw them
                self._efv_draw_overlays(self.efv_videoFrame, self.efv_overlays_data, self.efv_currentFrame)

            self.efv_showVideoFrame(self.efv_videoFrame, self.efv_videoView)

            self.efv_frameEdit.activateWindow()
            self.efv_frameEdit.setText(str(self.efv_currentFrame))
            self.efv_videoSlider.setValue(self.efv_currentFrame)

    def efv_updateModel(self, dataTable, mat, currentFrame):
        '''
        Function used to set up efViewer whenever a new event(log) is requested.
        :param dataTable: GUI component to be changed
        :param mat: input data for model
        :return:
        '''
        if currentFrame not in \
                self.efv_grb_index:
            self.statusBar().showMessage("Missing frame " + str(currentFrame), 1000)
            print("Missing frame " + str(currentFrame))
            return 0
        # header cleaning for resimed logs
        header = [h for h in self.efv_header if h[0] in mat.keys()]
        model = TableViewModel(self, mat, header, currentFrame)
        oldModel, oldSelectionModel = dataTable.model(), dataTable.selectionModel()
        if oldModel is not None and oldSelectionModel is not None:
            oldModel.deleteLater()
            oldSelectionModel.deleteLater()
        delegate = CustomDelegate(self)
        dataTable.setModel(model)
        dataTable.setItemDelegate(delegate)
        for idx in delegate.getComboIndexes(model):
            dataTable.openPersistentEditor(idx)
        dataTable.resizeRowsToContents()
        dataTable.setColumnWidth(0, dataTable.width() - dataTable.verticalHeader().width())
        dataTable.verticalHeader().setSectionsMovable(True)
        dataTable.verticalHeader().setDragEnabled(True)
        dataTable.verticalHeader().setDragDropMode(QtWidgets.QAbstractItemView.InternalMove)
        return 1

    def efv_generateVideoFrame(self, videoView, currentFrame):
        """
        Function checks if  colored video is requested, depending on that generates frame.
        If frame generation throws exception video view is cleared and error message is printed.
        :param videoView: GUI component to be changed
        :return: video frame
        """
        try:
            if self.efv_videoEnch.isVisible():
                self.efv_videoHandler.is_colored = self.efv_videoEnch.coloredBox.isChecked()
                videoFrame = self.efv_videoHandler.generate_frame(currentFrame,
                                                                  by_gid=True,
                                                                  enhancement=(
                                                                      self.efv_videoEnch.brightnessSlider.value(),
                                                                      self.efv_videoEnch.contrastSlider.value()))
            else:
                videoFrame = self.efv_videoHandler.generate_frame(currentFrame, by_gid=True)
            return videoFrame
        except Exception as e:
            videoView.scene().clear()
            print('Error generating frame', e)

    def efv_showVideoFrame(self, videoFrame, videoView):
        """
        Draw frame_img on QtGui.QImage.
        :param videoFrame: video frame to be displayed
        :param videoView: GUI component to be changed
        :return: None
        """
        videoFrame = cv2.cvtColor(videoFrame, cv2.COLOR_BGR2RGB)  # color mapping
        img = QtGui.QImage(videoFrame, videoFrame.shape[1], videoFrame.shape[0], QtGui.QImage.Format_RGB888)
        videoView.scene().clear()
        pix = QtGui.QPixmap.fromImage(img)
        pixmap = videoView.scene().addPixmap(pix)
        pixmap.setZValue(-1)
        videoView.setMinimumSize(1, 1)

    def efv_nextEvent(self):
        '''
        Function to show next event and save current event if saveAction is enabled.
        Signal to save resimed data is emitted if Resim Viewer is opened.
        :return: None
        '''
        self.efv_currentEvent = ((self.efv_currentEvent) + 1) % len(self.efv_events)
        self.efv_eventEdit.setText(str(self.efv_currentEvent + 1))
        if self.efv_saveAction.isChecked():
            self.efv_incrementSaveID()

            # signal emitted to resim viewer
            if self.efv_is2ndWindow:
                self.efv_saveSignal.emit()

        self.efv_updateEvent()

    def efv_prevEvent(self):
        '''
        Function to show previous event.
        :return: None
        '''
        self.efv_currentEvent = (self.efv_currentEvent - 1) % len(self.efv_events)
        self.efv_eventEdit.setText(str(self.efv_currentEvent + 1))
        self.efv_updateEvent()

    def efv_setEvent(self):
        '''
        Function to set event.
        :return: None
        '''
        self.setFocus()
        self.efv_currentEvent = (int(self.efv_eventEdit.text()) - 1) % len(self.efv_events)
        self.efv_eventEdit.setText(str(self.efv_currentEvent + 1))
        self.efv_updateEvent()

    def efv_getEventFiles(self, dest_folder, rtag=None):
        '''
        Get event files (.avi, .tavi, .mat, .dat). In case when rtag is specified function searches for resimed logs.
        :param dest_folder: folder with logs
        :param rtag: rtag of resimed logs
        :return: event files (.avi, .tavi, .mat, .dat)
        '''
        if rtag:
            matFilePath = os.path.join(dest_folder,
                                   self.efv_events[self.efv_currentEvent]['logName']) + rtag + ".mat"
            dest_folder = os.path.split(dest_folder)[0]
        else:
            matFilePath = os.path.join(dest_folder,
                                       self.efv_events[self.efv_currentEvent]['logName']) + ".mat"

        aviFilePath = os.path.join(dest_folder,
                                       self.efv_events[self.efv_currentEvent]['logName']) + ".avi"
        taviFilePath = os.path.join(dest_folder,
                                        self.efv_events[self.efv_currentEvent]['logName']) + ".tavi"
        datFilePath = os.path.join(dest_folder,
                                       self.efv_events[self.efv_currentEvent]['logName']) + ".dat"

        return matFilePath, aviFilePath, taviFilePath, datFilePath

    def efv_updateEvent(self, quietLoad=False, reload=False):
        """
        Function updates event. Check if current video name matched currently loaded .mat, if not - load new .mat file,
        set matUpperIndex and matLowerIndexand determine structure of mat in which grab_index is stored.
        Set text of efv_frameEdit and max/min values of VideoSlider
        :param quietLoad:
        :param reload:
        :return:
        """
        QtCore.QCoreApplication.processEvents()
        matFilePath, aviFilePath, taviFilePath, datFilePath = self.efv_getEventFiles(self.efv_loadDataLine.text())
        if os.path.isfile(matFilePath):
            needReload = self.efv_events[self.efv_currentEvent]['logName'] != os.path.splitext(
                         os.path.basename(self.efv_mat['__path__']))[0] or reload
            if needReload:
                if os.path.isfile(aviFilePath):
                    self.efv_videoHandler = vis2.Video(video_path=aviFilePath)
                elif os.path.isfile(taviFilePath):
                    self.efv_videoHandler = vis2.Video(video_path=taviFilePath)
                else:
                    self.efv_videoHandler = None
                    self.efv_videoView.scene().clear()
                    print('Missing video file!\nCannot open preview for current event')

                is_dat2p0 = self.efv_videoHandler.is_dat2p0
                if self.efv_loadBox.isChecked():
                    self.efv_mat = loadmat(matFilePath, variableName='mudp', sort=True, dat2p0=is_dat2p0)
                else:
                    self.efv_mat = loadmat(matFilePath, sort=True, dat2p0=is_dat2p0)

                if os.path.isfile(datFilePath):
                    # load GT
                    self.efv_load_groundtruth(datFilePath)

                # Determine where is grabIndex based on '__project__' value
                if self.efv_mat['__project__'] == 'DAT2.0':
                    self.is_dat2p0 = True
                    self.efv_grb_index = self.efv_mat['mudp']['vis']['vision_AEB_info']['visAEB']['imageIndex']
                else:
                    self.is_dat2p0 = False
                    self.efv_grb_index = self.efv_mat['mudp']['vis']['vision_traffic_sign_info']['imageIndex']

                self.matLowerIndex = self.efv_grb_index[0]
                self.matUpperIndex = self.efv_grb_index[-1]
                self.efv_overlays_activated = False  # Reset overlays - they will be initiated again

                # Set checkable options in right-click menu
                self._efv_setupMenu()

            # update mat for report generation purposes
            self.efv_mat.update({'saveID': self.saveID})
            self.fields_to_save = [['logName'], ['eventIndex'], ['saveID']]
            for field in self.fields_to_save:
                if field not in self.efv_header:
                    self.efv_header += [field]

            self.efv_mat.update(self.efv_events[self.efv_currentEvent])
            for key in deepcopy(self.efv_header):
                if key[-1].endswith('_corrected') or key[-1] == 'analysisStatus':
                    self.efv_header.pop(self.efv_header.index(key))
                    if not needReload:
                        mat_fields = reduce(operator.getitem, key[:-1], self.efv_mat)
                        mat_fields.pop(key[-1])
            if not quietLoad:
                self.efv_currentFrame = int(self.efv_events[self.efv_currentEvent]['eventIndex'])

                # Condition used when generating frames from fake report (if we have no EF report)
                if self.efv_currentFrame == -1:
                    self.efv_currentFrame = int(self.matLowerIndex)

            self.efv_frameEdit.setText(str(self.efv_currentFrame))
            self.efv_videoSlider.setMinimum(self.matLowerIndex)
            self.efv_videoSlider.setMaximum(self.matUpperIndex)
            self.efv_videoSlider.setValue(self.efv_currentFrame)
            self.efv_videoSlider.valueChanged.connect(partial(self.efv_updateFrameBySlider,
                                                              self.efv_frameEdit,
                                                              self.efv_updateFrame))

            self.efv_nextFrameBtn.setEnabled(True)
            self.efv_prevFrameBtn.setEnabled(True)
            self.efv_frameEdit.setEnabled(True)

            # signal emitted to resim viewer
            if self.efv_is2ndWindow:
                self.efv_eventSignal.emit(needReload)
        else:
            QtWidgets.QMessageBox.warning(self, 'Error', 'Missing data file!\nCannot open current event.')
            self.efv_videoHandler = None
            self.efv_nextFrameBtn.setEnabled(False)
            self.efv_prevFrameBtn.setEnabled(False)
            self.efv_frameEdit.setEnabled(False)
        self.efv_updateFrame()

    def efv_setCustomHeader(self):
        self.efv_customHeader = CustomHeader(self, self.efv_mat, self.efv_header)

    def efv_incrementSaveID(self):
        '''
        Function to increment save id and call self.efv_saveEvent()
        :return: None
        '''
        self.saveID += 1
        self.efv_mat.update({'saveID': self.saveID})
        self.efv_saveEvent(self.efv_dataTable, self.efv_videoFrame)

    def efv_saveEvent(self, dataTable, videoFrame, header=None, rtag=None):
        '''
        Function to save event from original and resimed logs.
        :param dataTable: GUI component which contains data to be saved
        :param videoFrame: video frame to be saved
        :param header: list of mat fields to be saved
        :param rtag: rtag of resimed log
        :return: None
        '''
        if not header:
            header = self.efv_header

        report_header = ['.'.join(h[-2:]) if (len(h) >= 2) else h[0] for h in header]
        data = [dataTable.model().getDataByHeader(h, save=True) for h in header]

        dataDict = dict(zip(report_header, data))
        dataDict['frame'] = videoFrame

        if rtag:
            reportData = str(self.efv_loadReportLine.text())[:-4] + rtag + '.report'
        else:
            reportData = str(self.efv_loadReportLine.text())[:-4] + '.report'

        with open(reportData, 'ab') as csvfile:
            pickle.dump(dataDict, csvfile)

        self.saveProject()
        self.statusBar().showMessage('Event saved!', 1000)

    def efv_getReportFiles(self):
        '''
        Function to get all report files from destination folder.
        :return:
        '''
        report_path = self.efv_loadReportLine.text()
        report_name = report_path.split('/')[-1]
        root = report_path.replace(report_name, '')
        report_name = report_name.replace('.csv', '')

        files = []
        for file in os.listdir(root):
            if file.startswith(report_name) and file.endswith(".report"):
                files += [os.path.join(root, file)]

        return files

    def efv_mergeReportFiles(self, files):
        '''
        Function to get data from report files. In case of several report files data is merged by saveID field.
        :param files: report files
        :return: data
        '''
        collected_data = []
        for reportData in files:
            data = []
            with open(reportData, 'rb') as reportDataFile:
                while True:
                    try:
                        data.append(pickle.load(reportDataFile))
                    except EOFError:
                        break
            resimRegex = re.search("(_rR\d{7,8})?(_rE\d{7,8})?_rM\d{7,8}(_rEfm\d{7,8})?", reportData)
            if resimRegex:
                rtag = resimRegex[0]
                for d in data:
                    keys_to_change = set(d.keys())
                    keys_to_change.remove('saveID')  # do not change ID
                    for k in keys_to_change:
                        new_key = k + rtag
                        d[new_key] = d.pop(k)

            collected_data += [data]

        # if only one report is in the root folder do not merge
        if len(collected_data) == 1:
            return collected_data[0]
        else:
            new_dict = defaultdict(dict)
            for d in collected_data:
                for elem in d:
                    new_dict[elem['saveID']].update(elem)  # merge dicts by unique ID
            merged_data = list(new_dict.values())
            return merged_data

    def efv_reportCsv(self):
        '''
        Function to create csv report. All data is saved except for images.
        :return: None
        '''
        reportOut = str(self.efv_loadReportLine.text())[:-4] + '_corrected.csv'

        files = self.efv_getReportFiles()
        data = self.efv_mergeReportFiles(files)

        header = ['logName',
                  'eventFinderID',
                  'eventID',
                  'eventComment',
                  'eventIndex',
                  'eventDuration',
                  'eventColumnID',
                  'ME_API',
                  'ME_SW',
                  'VFP_ver']

        header2 = set()
        frame_keys = []
        for d in data:
            header2.update(d.keys())
            frame_keys += [k for k in d.keys() if k.startswith('frame')]
        header += sorted(list(header2 - set(header) - set(frame_keys)))

        with open(reportOut, 'w', newline='') as reportOutFile:
            writer = csv.DictWriter(reportOutFile, header, dialect='excel', extrasaction='ignore')
            writer.writeheader()
            writer.writerows(data)

    def efv_reportGen(self):
        '''
        Function to create RVR/VTR reports.
        :return:
        '''
        reportDialog = ReportDialog(self)
        if reportDialog.exec_():
            reportInfo = reportDialog.getInfo()
            reportDialog.close()
        else:
            reportDialog.close()
            return

        status_color = {'Fixed': WD_COLOR_INDEX.BRIGHT_GREEN,
                         'Reproduced': WD_COLOR_INDEX.BRIGHT_GREEN,
                         'Improved': WD_COLOR_INDEX.YELLOW,
                         'Rejected': WD_COLOR_INDEX.TURQUOISE,
                         'Not fixed': WD_COLOR_INDEX.RED,
                         'Not reproduced': WD_COLOR_INDEX.RED,
                         'Unverified': WD_COLOR_INDEX.PINK,
                         'Degradation': WD_COLOR_INDEX.GRAY_50}

        files = self.efv_getReportFiles()
        data = self.efv_mergeReportFiles(files)

        if reportInfo['type'] == 'vtr':
            doc = docx.Document(os.path.join('data', 'VTR_FORD_YYYYMMDD_ADAS-XXX_rTAG'))
            doc.tables[0].cell(0, 1).text = reportInfo['ticketNo']  # JIRA ticket identifier
            doc.tables[0].cell(1, 1).text = reportInfo['ticketTitle']  # Title of ticket
            doc.tables[0].cell(2, 1).text = reportInfo['module']  # Problem module
            doc.tables[0].cell(3, 1).text = 'API' + reportInfo['api'] + ' SW' + reportInfo['soft']  # ME resim version
            doc.tables[0].cell(4, 1).text = reportInfo['vfp']  # VFP version
            doc.tables[0].cell(5, 1).text = reportInfo['fordtag']  # Ford Release Tag
            doc.tables[0].cell(6, 1).text = reportInfo['author']  # Author
            doc.tables[0].cell(7, 1).text = datetime.now().strftime('%d/%m/%Y')  # Date
            doc.tables[0].cell(9, 1).text = reportInfo['docs']  # Related Documents
            for i in range(1, 7):
                doc.tables[2].cell(i, 1).text = '/' + str(len(data))
        else:
            doc = docx.Document(os.path.join('data', 'RVR_FORD_YYYYMMDD_ADAS-XXX_rTAG'))
            doc.tables[0].cell(0, 1).text = reportInfo['ticketNo']  # JIRA ticket identifier
            doc.tables[0].cell(1, 1).text = reportInfo['ticketTitle']  # Title of ticket
            doc.tables[0].cell(2, 1).text = reportInfo['module']  # Problem module
            doc.tables[0].cell(3, 1).text = 'API' + reportInfo['api'] + ' SW' + reportInfo['soft']  # ME resim version
            doc.tables[0].cell(4, 1).text = reportInfo['vfp']  # VFP version
            doc.tables[0].cell(5, 1).text = reportInfo['author']  # Author
            doc.tables[0].cell(6, 1).text = datetime.now().strftime('%d/%m/%Y')  # Date
            doc.tables[0].cell(8, 1).text = reportInfo['docs']  # Related Documents
            for i in range(1, 3):
                doc.tables[2].cell(i, 1).text = '/' + str(len(data))

        doc.tables[1].style.font.size = Pt(10)
        for id, d in enumerate(data):
            # general
            cells = doc.tables[1].add_row().cells
            cells[0].text = str(id + 1) + '.'  # No
            if 'logName' in d.keys():
                cells[1].text = d['logName']  # LogName
            if 'eventIndex_corrected' in d.keys():
                cells[2].text = d['eventIndex_corrected']
            elif 'eventIndex' in d.keys():
                cells[2].text = d['eventIndex']  # Index
            if 'ME_API' in d.keys() and 'ME_SW' in d.keys():
                cells[3].text = 'API{} SW{}'.format(
                    d['ME_API'], d['ME_SW'])  # Orginal SW
            cells[4].text = 'Available'  # Availability
            if 'analysisStatus' in d.keys():
                run = cells[5].paragraphs[0].add_run(d['analysisStatus'])  # analysisStatus
                if d['analysisStatus'] in status_color.keys():
                    run.font.highlight_color = \
                        status_color[d['analysisStatus']]

            # detail
            table = doc.add_table(4, 3, doc.tables[1].style)
            sizes = [300000, 6000000, 6000000]
            for i in range(4):
                for j in range(3):
                    table.cell(i, j).width = sizes[j]
                if i > 0:
                    table.cell(i, 1).merge(table.cell(i, 2))

            table.cell(0, 0).text = str(id + 1) + '.'  # No
            if 'logName' in d.keys():
                table.cell(0, 1).text = d['logName']  # LogName
            if 'eventIndex_corrected' in d.keys():
                table.cell(0, 2).text = 'Grab Index: {}'.format(d['eventIndex_corrected'])
            elif 'eventIndex' in d.keys():
                table.cell(0, 2).text = 'Grab Index: {}'.format(d['eventIndex'])  # Index
            else:
                table.cell(0, 2).text = 'Grab Index: '
            table.cell(1, 1).paragraphs[0].add_run('Status: ')
            if 'analysisStatus' in d.keys():
                run = table.cell(1, 1).paragraphs[0].add_run('{}'.format(d['analysisStatus']))  # Status
                if d['analysisStatus'] in status_color.keys():
                    run.font.highlight_color = \
                        status_color[d['analysisStatus']]
            else:
                table.cell(1, 1).text = 'Status: '

            frame_keys = [k for k in d.keys() if k.startswith('frame')]
            if len(frame_keys) == 1:
                image_width = 7000000
            else:
                image_width = 4000000

            for f_key in frame_keys:
                frame = d[f_key]
                if len(frame.shape) == 3:
                    cv2.imwrite('temp.png', cv2.cvtColor(frame, cv2.COLOR_BGR2RGB))
                else:
                    cv2.imwrite('temp.png', cv2.cvtColor(frame, cv2.COLOR_GRAY2RGB))
                paragraph = table.cell(2, 1).paragraphs[0]
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                run = paragraph.add_run()
                run.add_picture('temp.png', width=image_width)
                os.remove('temp.png')  # Frame

            if 'eventComment_corrected' in d.keys():
                table.cell(3, 1).text = 'Comment: {}'.format(d['eventComment_corrected'])
            elif 'eventComment' in d.keys():
                table.cell(3, 1).text = 'Comment: {}'.format(d['eventComment'])  # Comment
            doc.add_page_break()

        dir = os.path.dirname(self.efv_loadReportLine.text())
        file = '{}_FORD_{}_{}_rM{}{}.docx'.format(
            reportInfo['type'].upper(),
            datetime.now().strftime('%Y%m%d'),
            reportInfo['ticketNo'],
            reportInfo['api'],
            reportInfo['soft'])
        doc.save(os.path.join(dir, file))

    def efv_openBindMenu(self):
        self.efv_bindMenu.setVisible(not self.efv_bindMenu.isVisible())

    def efv_enchVideo(self):
        self.efv_videoEnch.setVisible(not self.efv_videoEnch.isVisible())
        self.efv_videoEnch.move(self.mapFromGlobal(self.cursor().pos()))
        self.efv_updateFrame()

    def efv_customTableMenuRequested(self, pos):
        menu = QtWidgets.QMenu(self)
        menu.addAction(self.efv_customizeTableAction)
        menu.addAction(self.efv_bindMenuAction)
        menu.addAction(self.efv_saveAction)
        menu.addAction(self.efv_genCsv)
        menu.addAction(self.efv_genReport)
        menu.popup(self.efv_dataTable.viewport().mapToGlobal(pos))

    def efv_customFrameMenuRequested(self, pos):
        """
        Define how right-clicked menu will look like (Menu with video enhancement and overlays)
        :param pos: Mouse cursor position
        :return: None
        """
        menu = QtWidgets.QMenu(self)
        menu.addAction(self.efv_videoEnchAction)
        menu.addSeparator()
        menu.addAction(self.efv_aflLayoutAction)
        menu.addAction(self.efv_hrsLayoutAction)
        menu.addSeparator()
        menu.addAction(self.efv_tsrLayoutAction)
        menu.addAction(self.efv_tsrPlusLayoutAction)
        menu.addSeparator()
        menu.addAction(self.efv_obj2dLayoutAction)
        menu.addAction(self.efv_obj3dLayoutAction)
        menu.addSeparator()
        menu.addAction(self.efv_lksIndLayoutAction)
        menu.addAction(self.efv_lksParLayoutAction)
        menu.addAction(self.efv_lksHspLayoutAction)
        menu.addAction(self.efv_lksReLayoutAction)
        menu.addAction(self.efv_lksBarLayoutAction)
        menu.addSeparator()
        menu.addAction(self.efv_lksHostLayoutAction)
        menu.addAction(self.efv_lksNextLLayoutAction)
        menu.addAction(self.efv_lksNextRLayoutAction)
        menu.addAction(self.efv_lksBordLayoutAction)
        menu.addSeparator()
        menu.addAction(self.efv_fusLayoutAction)
        menu.addAction(self.efv_fusPedLayoutAction)
        menu.addAction(self.efv_tselLayoutAction)
        menu.addAction(self.efv_tselPathLayoutAction)
        menu.addAction(self.efv_pcaLayoutAction)
        menu.addSeparator()
        menu.addAction(self.efv_failSafeLayoutAction)

        menu.addSeparator()
        menu.addAction(self.efv_gtLayoutAction)
        menu.popup(self.efv_videoView.mapToGlobal(pos))

    def viewWheelEvent(self, event):
        zoomAcceleration = 1.05 if event.angleDelta().y() > 0 else 0.95
        self.efv_videoView.setTransformationAnchor(QtWidgets.QGraphicsView.AnchorUnderMouse)
        self.efv_videoView.scale(zoomAcceleration, zoomAcceleration)

    def _efv_setupMenu(self):
        """
        Set which items from right-click menu should be chackable depending on the project that is currently used
        :return: None
        """
        if self.is_dat2p0:
            for overlay in [self.efv_lksIndLayoutAction, self.efv_lksParLayoutAction, self.efv_lksHspLayoutAction,
                            self.efv_lksReLayoutAction, self.efv_lksBarLayoutAction, self.efv_failSafeLayoutAction,
                            self.efv_fusLayoutAction, self.efv_fusPedLayoutAction, #self.efv_tselLayoutAction,
                            self.efv_tselPathLayoutAction, self.efv_pcaLayoutAction]:
                overlay.setEnabled(False)
                overlay.setCheckable(False)
        else:
            for overlay in [self.efv_lksHostLayoutAction, self.efv_lksNextLLayoutAction, self.efv_lksNextRLayoutAction,
                            self.efv_lksBordLayoutAction, self.efv_obj3dLayoutAction]:
                overlay.setEnabled(False)
                overlay.setCheckable(False)


class CustomHeader(QtWidgets.QDialog, customHeader.Ui_Dialog):
    def __init__(self, parent, mat, oldHeader):
        super(self.__class__, self).__init__(parent)
        self.setupUi(self)
        self.setModal(True)

        self.mat = mat

        self.saveBtn.clicked.connect(self.save)

        self.show()

        self.fieldsTree.clear()
        self.fillTree(self.fieldsTree.invisibleRootItem(), self.mat, oldHeader)

    def fillTree(self, item, value, oldHeader):
        if type(value) is dict:
            for key, val in sorted(iter(value.items())):
                if not key.startswith('__'):
                    child = QtWidgets.QTreeWidgetItem()
                    child.setText(0, key)
                    item.addChild(child)
                    if type(val) is dict and key != 'groundtruth':
                        self.fillTree(child, val, oldHeader)
                    else:
                        child.setFlags(child.flags() | QtCore.Qt.ItemIsUserCheckable | QtCore.Qt.ItemIsSelectable)
                        childPath = self.getFullTreePath(child)
                        child.setCheckState(0, 2 * (childPath in oldHeader))

    def getFullTreePath(self, item):
        path = [str(item.text(0))]
        while item.parent():
            item = item.parent()
            path.append(str(item.text(0)))
        path.reverse()
        return path

    def getChecked(self, item, paths):
        if item.childCount():
            for i in range(item.childCount()):
                self.getChecked(item.child(i), paths)
        elif item.checkState(0):
            paths.append(self.getFullTreePath(item))

    def save(self):
        pathsToSave = []
        self.getChecked(self.fieldsTree.invisibleRootItem(), pathsToSave)
        self.parent().efv_header = deepcopy(pathsToSave)
        self.parent().efv_updateFrame()
        self.parent().efv_bindMenu.updateComboBox()
        self.close()

    def closeEvent(self, QCloseEvent):
        self.mat = None
        self.deleteLater()


class EnchancerWidget(QtWidgets.QWidget, videoEnchancer.Ui_Form):
    def __init__(self, parent):
        super(self.__class__, self).__init__(parent)
        self.setupUi(self)

        self.brightnessSlider.valueChanged.connect(self.parent().efv_updateFrame)
        self.contrastSlider.valueChanged.connect(self.parent().efv_updateFrame)
        self.coloredBox.stateChanged.connect(self.parent().efv_updateFrame)
        self.hide()

    def mousePressEvent(self, event):
        self.moveVector = self.pos() - event.globalPos()

    def mouseMoveEvent(self, event):
        if self.moveVector is not None:
            newPos = event.globalPos() + self.moveVector
            self.move(newPos)

    def mouseReleaseEvent(self, event):
        self.moveVector = None

    def closeEvent(self, event):
        self.deleteLater()


class BindDialog(QtWidgets.QDialog, bindDialog.Ui_bindDialog):
    def __init__(self, parent):
        super(self.__class__, self).__init__(parent)
        self.setupUi(self)
        self.setWindowFlags(QtCore.Qt.Tool)

        self.bind_1 = QtWidgets.QShortcut(QtGui.QKeySequence.fromString('1'), self.parent())
        self.bind_1.activated.connect(partial(self.activateBind, 1))
        self.bind_2 = QtWidgets.QShortcut(QtGui.QKeySequence.fromString('2'), self.parent())
        self.bind_2.activated.connect(partial(self.activateBind, 2))
        self.bind_3 = QtWidgets.QShortcut(QtGui.QKeySequence.fromString('3'), self.parent())
        self.bind_3.activated.connect(partial(self.activateBind, 3))
        self.bind_4 = QtWidgets.QShortcut(QtGui.QKeySequence.fromString('4'), self.parent())
        self.bind_4.activated.connect(partial(self.activateBind, 4))
        self.bind_5 = QtWidgets.QShortcut(QtGui.QKeySequence.fromString('5'), self.parent())
        self.bind_5.activated.connect(partial(self.activateBind, 5))
        self.bind_6 = QtWidgets.QShortcut(QtGui.QKeySequence.fromString('6'), self.parent())
        self.bind_6.activated.connect(partial(self.activateBind, 6))
        self.bind_7 = QtWidgets.QShortcut(QtGui.QKeySequence.fromString('7'), self.parent())
        self.bind_7.activated.connect(partial(self.activateBind, 7))
        self.bind_8 = QtWidgets.QShortcut(QtGui.QKeySequence.fromString('8'), self.parent())
        self.bind_8.activated.connect(partial(self.activateBind, 8))
        self.bind_9 = QtWidgets.QShortcut(QtGui.QKeySequence.fromString('9'), self.parent())
        self.bind_9.activated.connect(partial(self.activateBind, 9))

        self.updateComboBox()

        self.setVisible(False)

    def updateComboBox(self):
        text = str(self.fieldComboBox.currentText())
        self.fieldComboBox.clear()
        for h in self.parent().efv_header + [['analysisStatus']]:
            self.fieldComboBox.addItem(h[-1], h)
        if text in [self.fieldComboBox.itemText(i) for i in range(self.fieldComboBox.count())]:
            self.fieldComboBox.setCurrentText(text)

    def activateBind(self, val):
        header = str(self.fieldComboBox.currentText())
        fullHeader = self.fieldComboBox.currentData()

        text = eval('self.lineEdit_{}.text()'.format(val))
        if str(text) != '':
            if header == 'analysisStatus':
                if not fullHeader in self.parent().efv_header:
                    self.parent().efv_header.append(fullHeader)
                    self.parent().efv_mat['analysisStatus'] = text
            self.parent().efv_dataTable.model().setDataByHeader(self.fieldComboBox.currentData(), text)

            if self.checkBox.isChecked():
                self.parent().efv_nextEvent()
            else:
                self.parent().efv_updateFrame()

    def getBinds(self):
        binds = []
        for i in range(1,10):
            binds.append(eval('self.lineEdit_{}.text()'.format(i)))
        return binds

    def setBinds(self, binds):
        for i in range(1,10):
            binds.append(eval('self.lineEdit_{}.setText("{}")'.format(i, binds[i-1])))

    def closeEvent(self, event):
        self.setVisible(False)
        event.ignore()


class ReportDialog(QtWidgets.QDialog, reportGenWindow.Ui_Dialog):
    def __init__(self, parent):
        super(self.__class__, self).__init__(parent)
        self.setupUi(self)

        self.browseBtn.clicked.connect(partial(self.parent().loadFile, self.lineEdit_9))

    def getInfo(self):
        return {'type': 'vtr' if self.vtrBtn.isChecked() else 'rvr',
                'ticketNo': str(self.lineEdit_1.text()),
                'ticketTitle': str(self.lineEdit_2.text()),
                'module': str(self.lineEdit_3.text()),
                'api': str(self.lineEdit_4.text()),
                'soft': str(self.lineEdit_5.text()),
                'vfp': str(self.lineEdit_6.text()),
                'fordtag': str(self.lineEdit_7.text()),
                'author': str(self.lineEdit_8.text()),
                'docs': os.path.basename(str(self.lineEdit_9.text()))}

    def closeEvent(self, QCloseEvent):
        self.deleteLater()


class ResimViewer(QtWidgets.QDialog, resimViewer.Ui_Dialog):
    def __init__(self, parent):
        super(self.__class__, self).__init__(parent)
        self.setupUi(self)

        # variable initialization
        self.rv_mat = None
        self.rv_overlays_activated = False
        self.parent().efv_is2ndWindow = True

        # GUI initialization and signals connections
        self.splitter.setSizes([1, 600])
        self.splitter_2.setSizes([60, 1])
        self.splitter_2.splitterMoved.connect(self.rv_updateFrame)

        self.rv_dataTable.setContextMenuPolicy(QtCore.Qt.CustomContextMenu)
        self.rv_videoView.setContextMenuPolicy(QtCore.Qt.CustomContextMenu)
        scene = QtWidgets.QGraphicsScene(self)
        self.rv_videoView.setScene(scene)
        self.rv_videoView.setDragMode(QtWidgets.QGraphicsView.ScrollHandDrag)
        self.rv_videoView.wheelEvent = self.viewWheelEvent

        self.rv_loadFolderBtn.clicked.connect(partial(self.rv_loadFolder, self.parent().efv_loadDataLine.text()))
        self.parent().efv_eventSignal.connect(self.rv_updateEvent)
        self.parent().efv_frameSignal.connect(self.rv_updateFrame)
        self.parent().efv_saveSignal.connect(self.rv_saveEvent)
        self.parent().efv_saveFrame10ShortCut.activated.connect(self.rv_saveEvent)

        self.rv_synchronized.setChecked(True)
        self.rv_synchronized.stateChanged.connect(self.rv_enableAsync)

        self.rv_videoSlider.setEnabled(False)
        self.rv_frameEdit.setEnabled(False)
        self.rv_prevFrameBtn.setEnabled(False)
        self.rv_nextFrameBtn.setEnabled(False)

        self.rv_prevFrameBtn.clicked.connect(partial(self.parent().efv_prevFrame, 2, self.rv_frameEdit, self.rv_updateFrame))
        self.rv_nextFrameBtn.clicked.connect(partial(self.parent().efv_nextFrame, 2, self.rv_frameEdit, self.rv_updateFrame))

        self.rv_frameEdit.returnPressed.connect(partial(self.parent().efv_setFrame, self.rv_updateFrame))

    def rv_loadFolder(self, root_folder=None):
        '''
        Function to load resim folder and determine if rtag is valid.
        :param root_folder: root folder taken from event finder viewer.
        :return: None
        '''
        resim_root = QtWidgets.QFileDialog.getExistingDirectory(self, directory=root_folder,
                                                          options=QtWidgets.QFileDialog.ShowDirsOnly)
        if not resim_root:
            return

        rtag = resim_root.split('/')[-1]
        rtag = '_' + rtag
        resimRegex = re.search("(_rR\d{7,8})?(_rE\d{7,8})?_rM\d{7,8}(_rEfm\d{7,8})?", rtag)
        if resimRegex:
            self.rv_rtag = rtag
        else:
            QtWidgets.QMessageBox.warning(self, 'Error', 'Resim folder is not valid!\nCannot open current event.')
            return

        self.rv_loadDataLine.setText(resim_root)

        self.rv_updateEvent(True)

    def rv_enableAsync(self):
        '''
        Function to connect/disconnect signal from efv and change rv widgets status.
        Example: When rv_synchronized checkbox is set to true method connects signal from event viewer class
        to rv_updateFrame method.
        Widgets from resim viewer GUI which are used to change frames are disabled.
        :return: None
        '''
        syncChecked = self.rv_synchronized.isChecked()

        if syncChecked:
            self.parent().efv_frameSignal.connect(self.rv_updateFrame)
        else:
            self.parent().efv_frameSignal.disconnect(self.rv_updateFrame)

        self.rv_videoSlider.setEnabled(not syncChecked)
        self.rv_frameEdit.setEnabled(not syncChecked)
        self.rv_prevFrameBtn.setEnabled(not syncChecked)
        self.rv_nextFrameBtn.setEnabled(not syncChecked)

    def rv_updateFrame(self):
        '''
        Function to update frame in resim viewer.
        If synchronized mode is checked then video frame is copied from event finder viewer. Otherwise new frame is generated.
        Data table and overlays are also updated if needed.
        :return: None
        '''
        if not self.rv_mat:
            return

        syncChecked = self.rv_synchronized.isChecked()

        if syncChecked:
            self.rv_currentFrame = self.parent().efv_currentFrame
            self.rv_frameEdit.setText(str(self.rv_currentFrame))

            self.rv_videoFrame = deepcopy(self.parent().efv_videoFrame)
        else:
            self.rv_currentFrame = int(self.rv_frameEdit.text())

            if self.parent().efv_videoHandler is not None:
                self.rv_videoFrame = self.parent().efv_generateVideoFrame(self.rv_videoView, self.rv_currentFrame)

        self.rv_videoSlider.setValue(self.rv_currentFrame)

        try:
            if not self.parent().efv_updateModel(self.rv_dataTable, self.rv_mat, self.rv_currentFrame):
                return
        except IndexError:
            print('Missing frame in resimulated log.')

        # case when resim viewer is initalized after activation of overlays on original frame
        if self.parent().efv_overlays_activated and not self.rv_overlays_activated:
            self.rv_videoFrame = self.parent().efv_generateVideoFrame(self.rv_videoView, self.rv_currentFrame)
            self.rv_overlays_data = self.parent()._efv_initiate_overlays(self.rv_videoFrame, self.rv_mat, self.rv_overlays_activated)

        self.rv_overlays_activated = self.parent().efv_overlays_activated

        if self.rv_overlays_activated:
            self.parent()._efv_draw_overlays(self.rv_videoFrame, self.rv_overlays_data, self.rv_currentFrame)

        self.parent().efv_showVideoFrame(self.rv_videoFrame, self.rv_videoView)

    def rv_updateEvent(self, needReload):
        '''
        Function to update event. If needReload is true then new mat is loaded.
        :param needReload: bool to indentify if new mat is needed
        :return: None
        '''
        if needReload:
            try:
                matFilePath, _, _, _ = self.parent().efv_getEventFiles(self.rv_loadDataLine.text(), self.rv_rtag)

                if self.parent().efv_loadBox.isChecked():
                    self.rv_mat = loadmat(matFilePath, variableName='mudp', sort=True, dat2p0=self.parent().is_dat2p0)
                else:
                    self.rv_mat = loadmat(matFilePath, sort=True, dat2p0=self.parent().is_dat2p0)

                self.rv_videoSlider.setMinimum(self.parent().matLowerIndex)
                self.rv_videoSlider.setMaximum(self.parent().matUpperIndex)
                self.rv_videoSlider.valueChanged.connect(partial(self.parent().efv_updateFrameBySlider,
                                                                 self.rv_frameEdit,
                                                                 self.rv_updateFrame))

                self.rv_overlays_activated = False

            except FileNotFoundError:
                QtWidgets.QMessageBox.warning(self, 'Error', ''.join([matFilePath, '\n', 'File not found!\nCannot open current event.']))
                return

        self.rv_mat.update({'saveID': self.parent().saveID})

        self.rv_currentFrame = int(self.parent().efv_frameEdit.text())

        self.rv_frameEdit.setText(str(self.rv_currentFrame))
        self.rv_videoSlider.setValue(self.rv_currentFrame)

        self.rv_synchronized.setChecked(True)

        self.rv_updateFrame()

    def rv_saveEvent(self):
        '''
        Function to update saveID, clean resim header and call saveEvent function.
        :return: None
        '''
        self.rv_mat.update({'saveID': self.parent().saveID})
        self.rv_header = [h for h in self.parent().efv_header if h[0] in self.rv_mat.keys()]

        self.parent().efv_saveEvent(self.rv_dataTable, self.rv_videoFrame, self.rv_header, self.rv_rtag)

    def viewWheelEvent(self, event):
        zoomAcceleration = 1.05 if event.angleDelta().y() > 0 else 0.95
        self.rv_videoView.setTransformationAnchor(QtWidgets.QGraphicsView.AnchorUnderMouse)
        self.rv_videoView.scale(zoomAcceleration, zoomAcceleration)

    def closeEvent(self, QCloseEvent):
        self.parent().efv_is2ndWindow = False
        self.deleteLater()
