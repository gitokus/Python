import datetime
import os
import sys
import docx

from PyQt5.QtWidgets import QApplication, QMainWindow, QFileDialog, QMessageBox
from PyQt5.uic import loadUi


class ReportGenerator(QMainWindow):
    def __init__(self):
        super(ReportGenerator, self).__init__()
        loadUi('mw.ui',self)

        self.lineEdit_output.setText(os.getcwd())
        self.generatBtn.clicked.connect(self.generate)
        self.docBtn.clicked.connect(self.loadDoc)
        self.outputBtn.clicked.connect(self.loadOutput)
        self.taviBtn.clicked.connect(self.loadTavi)
        self.tabWidget.currentChanged.connect(self.tabChanged)


    def loadDoc(self):
        path = '/'.join(str(self.lineEdit_15.text()).split('/')[:-1])
        if path == '':
            path = os.getcwd()
        file = QFileDialog.getOpenFileName(self, 'Open', path)[0]
        file = file.replace('\\\\', '\\') 
        file = file.replace('\\', '/')
        if file:
            file = str(file).split('/')[-1]
        else:
            file = ''
        self.lineEdit_15.setText(file)


    def loadOutput(self):
        path = '/'.join(str(self.lineEdit_output.text()).split('/')[:-1])
        if path == '':
            path = os.getcwd()
        file = QFileDialog.getExistingDirectory(self, 'Open', directory=path, options=QFileDialog.ShowDirsOnly)
        file = file.replace('\\\\','\\').replace('\\', '/')
        if not file:
            return
        self.lineEdit_output.setText(file)


    def loadTavi(self):
        path = '/'.join(str(self.lineEdit_tavi.text()).split('/')[:-1])
        if path == '':
            path = os.getcwd()
        file = QFileDialog.getExistingDirectory(self, directory=path,
               options=QFileDialog.ShowDirsOnly).replace('\\\\','\\').replace('\\', '/')
        if file:
            self.lineEdit_tavi.setText(file)
        else:
            return


    def generate(self):
        # try:
        document = docx.Document(r'templates\VTR_FORD_YYYYMMDD_ADAS-XXX_rTAG')

        self.getInfo()
        self.fillGeneralInfo(document)
        self.fillLogList(document)
        self.fillDetail(document)

        name = '_'.join(['VTR', 'FORD', datetime.datetime.now().strftime('%Y%m%d'), str(self.lineEdit_1.text()),
                         'rM' + str(self.lineEdit_4.text()) + str(self.lineEdit_5.text()) + '.docx'])
        document.save(os.path.join(str(self.lineEdit_output.text()), name))
        QMessageBox.information(self, 'Success', 'Report genereted successfully!')
        # except:
        #     QtGui.QMessageBox.warning(self, 'Warning', 'Error while generating report!')

    def getInfo(self):
        names = str(self.logsEdit.toPlainText()).splitlines()
        indexes = str(self.indexEdit.toPlainText()).splitlines()
        sw = str(self.swEdit.toPlainText()).splitlines()

        while len(indexes)<len(names):
            indexes.append('')
        while len(sw) < len(names):
            sw.append('')

        self.info = zip(names, indexes, sw)


    def fillGeneralInfo(self, doc):
        doc.tables[0].cell(0, 1).text = str(self.lineEdit_1.text())                         # JIRA ticket identifier
        doc.tables[0].cell(1, 1).text = str(self.lineEdit_2.text())                         # Title of ticket
        doc.tables[0].cell(2, 1).text = str(self.lineEdit_3.text())                         # Problem module
        doc.tables[0].cell(3, 1).text = \
            'API' + str(self.lineEdit_4.text()) + ' SW' + str(self.lineEdit_5.text())       # ME resim version
        doc.tables[0].cell(4, 1).text = str(self.lineEdit_6.text())                         # VFP version
        doc.tables[0].cell(5, 1).text = str(self.lineEdit_7.text())                         # Ford Release Tag
        doc.tables[0].cell(6, 1).text = str(self.lineEdit_8.text())                         # Author
        doc.tables[0].cell(7, 1).text = datetime.datetime.now().strftime('%d/%m/%Y')        # Date
        doc.tables[0].cell(8, 1).paragraphs[2].add_run(' ' + str(self.lineEdit_9.text()))
        doc.tables[0].cell(8, 1).paragraphs[3].add_run(' ' + str(self.lineEdit_10.text()))
        doc.tables[0].cell(8, 1).paragraphs[4].add_run(' ' + str(self.lineEdit_11.text()))
        doc.tables[0].cell(8, 1).paragraphs[6].add_run(' ' + str(self.lineEdit_12.text()))
        doc.tables[0].cell(8, 1).paragraphs[7].add_run(' ' + str(self.lineEdit_13.text()))  # Description
        doc.tables[0].cell(9, 1).text = str(self.lineEdit_15.text())                        # Related Documents


    def fillLogList(self, doc):
        for i, (aLogName, aIndex, aSw) in enumerate( self.info ): #enumerate over zip object with index
            cells = doc.tables[1].add_row().cells
            cells[0].text = str(i+1)+'.'   # No
            cells[1].text = aLogName       # LogName
            cells[2].text = aIndex         # Index
            cells[3].text = aSw            # Orginal SW
            cells[4].text = 'Available'           # Availability

        for i in range(2,8):
            doc.tables[2].cell(i,1).text = '/' + str(len(list(self.info)))

        doc.tables[1].style.font.size = docx.shared.Pt(10)

    def fillDetail(self, doc):
        def newTable(no):
            table = doc.add_table(4, 3, doc.tables[1].style)
            sizes = [300000, 6000000, 6000000]
            for i in range(4):
                for j in range(3):
                    table.cell(i, j).width = sizes[j]
                if i > 0:
                    table.cell(i,1).merge(table.cell(i, 2))

            table.cell(0, 0).text = str(no+1)+'.'                       # No
            table.cell(0, 1).text = list(self.info)[no][0]                    # LogName
            table.cell(0, 2).text = 'Grab Index: ' + list(self.info)[no][1]   # Index
            table.cell(1, 1).text = 'Status: '                          # Status
            if not str(self.lineEdit_tavi.text()) == '':
                path = str(self.lineEdit_tavi.text())
                if os.path.isfile(os.path.join(path, self.info[no][0] + '.tavi')) and \
                        os.path.isfile(os.path.join(path, self.info[no][0] + '.mat')):
                #     vH = vis2.videoHandler(os.path.join(path, self.info[no][0] + '.tavi', ),
                #                            os.path.join(path, self.info[no][0] + '.mat'))
                #     vH.showLayer('tsr', True)
                #     frame = vH.generateFrames(int(self.info[no][1]), fromData=True)[0]
                #     cv2.imwrite('temp.png', frame)
                    paragraph = table.cell(2, 1).paragraphs[0]
                    paragraph.alignment = 1
                    run = paragraph.add_run()
                #     run.add_picture('temp.png', width=7000000)
                #     print('hehe')
                #     vH.close()
                #     os.remove('temp.png')
            table.cell(3, 1).text = 'Comment: '                         # Comment
            doc.add_page_break()

        for i in enumerate(self.info):
            newTable(i)


    def tabChanged(self):
        if self.tabWidget.currentIndex() == 0:
            self.resize(410, 540)
        if self.tabWidget.currentIndex() == 1:
            self.resize(820, 540)

app=QApplication(sys.argv)
widget=ReportGenerator()
widget.show()
sys.exit(app.exec_())
